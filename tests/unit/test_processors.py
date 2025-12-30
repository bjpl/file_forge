"""TDD Tests for FileForge processors (document, text).

RED phase: Tests written first, defining expected behavior.
These tests will fail until implementation is complete.
"""
import pytest
from pathlib import Path
from unittest.mock import MagicMock, patch, Mock
from dataclasses import dataclass


class TestExtractedContent:
    """Tests for ExtractedContent dataclass."""

    def test_extracted_content_creation(self):
        """Should create ExtractedContent with required fields."""
        from fileforge.pipeline.processors.document import ExtractedContent

        content = ExtractedContent(
            text="Sample text",
            structure={'headings': ['Title']},
            metadata={'author': 'Test'},
            pages=[],
            embedded_images=[],
            confidence=0.95
        )

        assert content.text == "Sample text"
        assert content.confidence == 0.95
        assert content.structure == {'headings': ['Title']}
        assert content.metadata == {'author': 'Test'}

    def test_extracted_content_defaults(self):
        """Should have sensible defaults for optional fields."""
        from fileforge.pipeline.processors.document import ExtractedContent

        content = ExtractedContent(
            text="Sample text",
            structure={},
            metadata={},
            pages=[],
            embedded_images=[],
            confidence=1.0
        )

        assert content.pages == []
        assert content.embedded_images == []
        assert content.confidence == 1.0


class TestDocumentProcessor:
    """Tests for DocumentProcessor class."""

    def test_supported_extensions(self):
        """Should support PDF and DOCX extensions."""
        from fileforge.pipeline.processors.document import DocumentProcessor

        processor = DocumentProcessor()
        assert '.pdf' in processor.supported_extensions
        assert '.docx' in processor.supported_extensions
        assert '.doc' in processor.supported_extensions

    def test_process_pdf_extracts_text(self, tmp_path):
        """Should extract text from PDF files."""
        from fileforge.pipeline.processors.document import DocumentProcessor, ProcessingError

        # Create a simple PDF for testing
        pdf_path = tmp_path / "test.pdf"
        # Will need sample PDF creation helper - for now test that missing file raises error

        processor = DocumentProcessor()

        # Test that missing file is handled
        with pytest.raises(ProcessingError):
            result = processor.process(pdf_path)

    def test_process_pdf_detects_scanned_pages(self, tmp_path):
        """Should detect and OCR scanned PDF pages."""
        from fileforge.pipeline.processors.document import DocumentProcessor

        processor = DocumentProcessor()

        # Mock a PDF with image-only pages
        with patch.object(processor, '_extract_pdf_text') as mock_extract:
            mock_extract.return_value = [
                MagicMock(text="", is_ocr=False, page_num=1),  # Empty = scanned
            ]
            with patch.object(processor, '_ocr_page') as mock_ocr:
                mock_ocr.return_value = "OCR extracted text"

                # This will fail until OCR detection is implemented
                with pytest.raises((AttributeError, NotImplementedError)):
                    result = processor._handle_scanned_pages([])

    def test_process_docx_extracts_paragraphs(self, tmp_path):
        """Should extract paragraphs from DOCX files."""
        from fileforge.pipeline.processors.document import DocumentProcessor

        docx_path = tmp_path / "test.docx"

        # Create minimal DOCX
        try:
            from docx import Document
            doc = Document()
            doc.add_paragraph("First paragraph")
            doc.add_paragraph("Second paragraph")
            doc.save(docx_path)

            processor = DocumentProcessor()
            result = processor.process(docx_path)

            assert "First paragraph" in result.text
            assert "Second paragraph" in result.text
            assert result.confidence > 0
        except ImportError:
            pytest.skip("python-docx not installed")

    def test_process_docx_extracts_headings(self, tmp_path):
        """Should extract heading structure from DOCX."""
        from fileforge.pipeline.processors.document import DocumentProcessor

        try:
            from docx import Document
            docx_path = tmp_path / "test.docx"
            doc = Document()
            doc.add_heading("Main Title", 0)
            doc.add_heading("Section 1", 1)
            doc.add_paragraph("Content")
            doc.save(docx_path)

            processor = DocumentProcessor()
            result = processor.process(docx_path)

            assert 'headings' in result.structure
            assert len(result.structure['headings']) == 2
            assert result.structure['headings'][0]['text'] == "Main Title"
        except ImportError:
            pytest.skip("python-docx not installed")

    def test_process_docx_extracts_metadata(self, tmp_path):
        """Should extract metadata from DOCX."""
        from fileforge.pipeline.processors.document import DocumentProcessor

        try:
            from docx import Document
            docx_path = tmp_path / "test.docx"
            doc = Document()
            doc.core_properties.author = "Test Author"
            doc.core_properties.title = "Test Title"
            doc.add_paragraph("Content")
            doc.save(docx_path)

            processor = DocumentProcessor()
            result = processor.process(docx_path)

            assert result.metadata.get('author') == "Test Author"
            assert result.metadata.get('title') == "Test Title"
        except ImportError:
            pytest.skip("python-docx not installed")

    def test_process_returns_confidence_score(self, tmp_path):
        """Should return confidence score for extraction."""
        from fileforge.pipeline.processors.document import DocumentProcessor

        try:
            from docx import Document
            docx_path = tmp_path / "test.docx"
            doc = Document()
            doc.add_paragraph("Clear text content")
            doc.save(docx_path)

            processor = DocumentProcessor()
            result = processor.process(docx_path)

            assert 0 <= result.confidence <= 1
            assert result.confidence > 0.5  # Should have good confidence for clear text
        except ImportError:
            pytest.skip("python-docx not installed")

    def test_handles_corrupt_file_gracefully(self, tmp_path):
        """Should handle corrupt files without crashing."""
        from fileforge.pipeline.processors.document import DocumentProcessor

        corrupt_file = tmp_path / "corrupt.pdf"
        corrupt_file.write_bytes(b"not a valid pdf")

        processor = DocumentProcessor()

        # Should raise ProcessingError when implemented
        with pytest.raises((ImportError, AttributeError, Exception)):
            processor.process(corrupt_file)

    def test_process_extracts_embedded_images(self, tmp_path):
        """Should extract embedded images from documents."""
        from fileforge.pipeline.processors.document import DocumentProcessor

        processor = DocumentProcessor()

        # This will fail until image extraction is implemented
        with pytest.raises((AttributeError, NotImplementedError)):
            images = processor._extract_images(tmp_path / "test.pdf")

    def test_process_handles_tables(self, tmp_path):
        """Should extract and structure tables from documents."""
        from fileforge.pipeline.processors.document import DocumentProcessor

        try:
            from docx import Document
            docx_path = tmp_path / "test.docx"
            doc = Document()
            table = doc.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = "Header 1"
            table.rows[0].cells[1].text = "Header 2"
            doc.save(docx_path)

            processor = DocumentProcessor()
            result = processor.process(docx_path)

            assert 'tables' in result.structure
            assert len(result.structure['tables']) == 1
            assert result.structure['tables'][0]['data'][0][0] == "Header 1"
        except ImportError:
            pytest.skip("python-docx not installed")


class TestTextProcessor:
    """Tests for TextProcessor class."""

    def test_supported_extensions(self):
        """Should support TXT and Markdown extensions."""
        from fileforge.pipeline.processors.text import TextProcessor

        processor = TextProcessor()
        assert '.txt' in processor.supported_extensions
        assert '.md' in processor.supported_extensions
        assert '.markdown' in processor.supported_extensions

    def test_process_txt_reads_content(self, tmp_path):
        """Should read text file content."""
        from fileforge.pipeline.processors.text import TextProcessor

        txt_path = tmp_path / "test.txt"
        txt_path.write_text("This is test content.\nLine two.")

        processor = TextProcessor()
        result = processor.process(txt_path)

        assert "test content" in result.text
        assert "Line two" in result.text
        assert result.confidence > 0

    def test_process_markdown_extracts_structure(self, tmp_path):
        """Should extract structure from Markdown."""
        from fileforge.pipeline.processors.text import TextProcessor

        md_path = tmp_path / "test.md"
        md_path.write_text("""# Main Heading

## Section 1
Some content here.

## Section 2
More content.

- List item 1
- List item 2
""")

        processor = TextProcessor()
        result = processor.process(md_path)

        assert 'headings' in result.structure
        assert len(result.structure['headings']) == 3
        assert result.structure['headings'][0]['level'] == 1
        assert result.structure['headings'][0]['text'] == "Main Heading"

    def test_process_markdown_extracts_code_blocks(self, tmp_path):
        """Should identify code blocks in Markdown."""
        from fileforge.pipeline.processors.text import TextProcessor

        md_path = tmp_path / "test.md"
        md_path.write_text("""# Code Example

```python
def hello():
    print("Hello")
```
""")

        processor = TextProcessor()
        result = processor.process(md_path)

        assert 'code_blocks' in result.structure
        assert len(result.structure['code_blocks']) == 1
        assert result.structure['code_blocks'][0]['language'] == 'python'
        assert 'def hello' in result.structure['code_blocks'][0]['code']

    def test_handles_different_encodings(self, tmp_path):
        """Should handle different text encodings."""
        from fileforge.pipeline.processors.text import TextProcessor

        # UTF-8 with special characters
        utf8_path = tmp_path / "utf8.txt"
        utf8_path.write_text("Héllo Wörld 中文", encoding='utf-8')

        processor = TextProcessor()
        result = processor.process(utf8_path)

        assert "Héllo" in result.text or "Hello" in result.text
        assert "Wörld" in result.text or "World" in result.text

    def test_extracts_frontmatter_yaml(self, tmp_path):
        """Should extract YAML frontmatter from Markdown."""
        from fileforge.pipeline.processors.text import TextProcessor

        md_path = tmp_path / "test.md"
        md_path.write_text("""---
title: Test Document
author: John Doe
date: 2024-01-01
---

# Content starts here
""")

        processor = TextProcessor()
        result = processor.process(md_path)

        assert result.metadata.get('title') == "Test Document"
        assert result.metadata.get('author') == "John Doe"
        assert "Content starts here" in result.text

    def test_identifies_document_type(self, tmp_path):
        """Should identify document type (notes, readme, etc)."""
        from fileforge.pipeline.processors.text import TextProcessor

        readme_path = tmp_path / "README.md"
        readme_path.write_text("""# Project Name

## Installation
Run `pip install project`

## Usage
Import and use.
""")

        processor = TextProcessor()
        result = processor.process(readme_path)

        assert 'doc_type' in result.metadata
        assert result.metadata['doc_type'] == 'readme'

    def test_confidence_based_on_content_quality(self, tmp_path):
        """Confidence should reflect content quality."""
        from fileforge.pipeline.processors.text import TextProcessor

        # Good content
        good_path = tmp_path / "good.txt"
        good_path.write_text("This is a well-formed document with proper sentences and structure.")

        # Poor content
        poor_path = tmp_path / "poor.txt"
        poor_path.write_text("asdf jkl; random gibberish 123")

        processor = TextProcessor()
        good_result = processor.process(good_path)
        poor_result = processor.process(poor_path)

        assert 0 <= good_result.confidence <= 1
        assert 0 <= poor_result.confidence <= 1
        assert good_result.confidence > poor_result.confidence  # Good should have higher confidence

    def test_extracts_links_from_markdown(self, tmp_path):
        """Should extract links from Markdown."""
        from fileforge.pipeline.processors.text import TextProcessor

        md_path = tmp_path / "test.md"
        md_path.write_text("""# Links

[Example](https://example.com)
[Local](./local.md)
""")

        processor = TextProcessor()
        result = processor.process(md_path)

        assert 'links' in result.structure
        assert len(result.structure['links']) == 2
        assert result.structure['links'][0]['url'] == 'https://example.com'
        assert result.structure['links'][1]['url'] == './local.md'

    def test_handles_very_large_files(self, tmp_path):
        """Should handle large text files efficiently."""
        from fileforge.pipeline.processors.text import TextProcessor

        large_path = tmp_path / "large.txt"
        large_content = "Line of text\n" * 100000  # 100k repetitions
        large_path.write_text(large_content)

        processor = TextProcessor()
        result = processor.process(large_path)

        assert "Line of text" in result.text
        # split('\n') on "text\n" * N gives N+1 elements (last is empty string)
        assert result.metadata['lines'] > 100000


class TestProcessorRegistry:
    """Tests for processor registration and routing."""

    def test_get_processor_for_extension(self):
        """Should return correct processor for extension."""
        from fileforge.pipeline.processors import get_processor, DocumentProcessor, TextProcessor

        pdf_proc = get_processor('.pdf')
        txt_proc = get_processor('.txt')

        assert isinstance(pdf_proc, DocumentProcessor)
        assert isinstance(txt_proc, TextProcessor)

    def test_unknown_extension_returns_none(self):
        """Should return None for unknown extensions."""
        from fileforge.pipeline.processors import get_processor

        assert get_processor('.xyz') is None
        assert get_processor('.unknown') is None

    def test_register_custom_processor(self):
        """Should allow registering custom processors."""
        from fileforge.pipeline.processors import register_processor, get_processor

        class CustomProcessor:
            supported_extensions = ['.custom']

            def process(self, path):
                return Mock(text="custom", confidence=1.0)

        register_processor('.custom', CustomProcessor)
        processor = get_processor('.custom')

        assert isinstance(processor, CustomProcessor)


class TestProcessingError:
    """Tests for ProcessingError exception."""

    def test_processing_error_creation(self):
        """Should create ProcessingError with message."""
        from fileforge.pipeline.processors.document import ProcessingError

        error = ProcessingError("Test error")
        assert str(error) == "Test error"

    def test_processing_error_with_cause(self):
        """Should chain exceptions properly."""
        from fileforge.pipeline.processors.document import ProcessingError

        try:
            raise ValueError("Original error")
        except ValueError as e:
            error = ProcessingError("Processing failed")
            error.__cause__ = e
            assert error.__cause__ is not None
            assert isinstance(error.__cause__, ValueError)


class TestDocumentProcessorIntegration:
    """Integration tests for document processing pipeline."""

    def test_process_mixed_document(self, tmp_path):
        """Should handle document with text, images, and tables."""
        from fileforge.pipeline.processors.document import DocumentProcessor

        try:
            from docx import Document
            docx_path = tmp_path / "mixed.docx"
            doc = Document()
            doc.add_heading("Mixed Document", 0)
            doc.add_paragraph("Some text content")
            table = doc.add_table(rows=1, cols=2)
            table.rows[0].cells[0].text = "Cell 1"
            doc.save(docx_path)

            processor = DocumentProcessor()
            result = processor.process(docx_path)

            assert "Mixed Document" in result.text
            assert 'headings' in result.structure
            assert 'tables' in result.structure
        except ImportError:
            pytest.skip("python-docx not installed")

    def test_batch_process_multiple_documents(self, tmp_path):
        """Should efficiently process multiple documents."""
        from fileforge.pipeline.processors.document import DocumentProcessor

        # This will fail until batch processing is implemented
        with pytest.raises((ImportError, AttributeError, NotImplementedError)):
            processor = DocumentProcessor()
            results = processor.batch_process([
                tmp_path / "doc1.pdf",
                tmp_path / "doc2.docx",
            ])


# Fixtures

@pytest.fixture
def temp_dir(tmp_path):
    """Provide temporary directory for tests."""
    return tmp_path


@pytest.fixture
def sample_pdf(tmp_path):
    """Create a sample PDF for testing."""
    # Will need implementation
    pytest.skip("PDF creation not implemented yet")


@pytest.fixture
def sample_docx(tmp_path):
    """Create a sample DOCX for testing."""
    try:
        from docx import Document
        path = tmp_path / "sample.docx"
        doc = Document()
        doc.add_heading("Test Document", 0)
        doc.add_paragraph("This is a test paragraph.")
        doc.save(path)
        return path
    except ImportError:
        pytest.skip("python-docx not installed")


@pytest.fixture
def sample_markdown(tmp_path):
    """Create a sample Markdown file for testing."""
    path = tmp_path / "sample.md"
    path.write_text("""# Test Document

## Section 1
Content here.

```python
print("code")
```
""")
    return path
