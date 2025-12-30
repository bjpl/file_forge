"""Fixtures for creating sample files for testing.

These helpers create various file types for testing processors.
"""
from pathlib import Path
from typing import Optional


def create_sample_pdf(path: Path, content: str = "Sample PDF content") -> Path:
    """Create a simple PDF file for testing.

    Args:
        path: Path where PDF should be created
        content: Text content to include

    Returns:
        Path to created PDF

    Note:
        This is a placeholder. Real implementation would use reportlab or similar.
    """
    # Placeholder - will need real PDF generation
    path.write_bytes(b"%PDF-1.4\n")  # Minimal PDF header
    return path


def create_sample_docx(
    path: Path,
    paragraphs: list[str] = None,
    headings: list[tuple[str, int]] = None
) -> Path:
    """Create a DOCX file for testing.

    Args:
        path: Path where DOCX should be created
        paragraphs: List of paragraph texts
        headings: List of (text, level) tuples

    Returns:
        Path to created DOCX
    """
    try:
        from docx import Document

        doc = Document()

        if headings:
            for text, level in headings:
                doc.add_heading(text, level)

        if paragraphs:
            for para in paragraphs:
                doc.add_paragraph(para)
        elif not headings:
            doc.add_paragraph("Default test content")

        doc.save(path)
        return path
    except ImportError:
        raise ImportError("python-docx required for creating test DOCX files")


def create_sample_markdown(
    path: Path,
    content: Optional[str] = None
) -> Path:
    """Create a Markdown file for testing.

    Args:
        path: Path where Markdown should be created
        content: Markdown content (uses default if None)

    Returns:
        Path to created Markdown file
    """
    if content is None:
        content = """# Test Document

## Introduction
This is a test markdown document.

## Features
- Feature 1
- Feature 2

```python
# Code example
def test():
    pass
```
"""

    path.write_text(content)
    return path


def create_sample_text(path: Path, content: str = "Sample text content") -> Path:
    """Create a plain text file for testing.

    Args:
        path: Path where text file should be created
        content: Text content

    Returns:
        Path to created text file
    """
    path.write_text(content)
    return path


def create_sample_image(
    path: Path,
    width: int = 640,
    height: int = 480,
    color: tuple = (128, 128, 128)
) -> Path:
    """Create a sample image file for testing.

    Args:
        path: Path where image should be created
        width: Image width in pixels
        height: Image height in pixels
        color: RGB color tuple (default: gray)

    Returns:
        Path to created image file
    """
    try:
        from PIL import Image

        # Create image with specified dimensions and color
        img = Image.new('RGB', (width, height), color)
        img.save(path)
        return path
    except ImportError:
        raise ImportError("Pillow required for creating test image files")
