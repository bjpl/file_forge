"""Document processor for PDF and DOCX files.

Handles structured document extraction including text, metadata, images, and tables.
"""
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from docx import Document
except ImportError:
    Document = None


class ProcessingError(Exception):
    """Exception raised when document processing fails."""
    pass


@dataclass
class ExtractedContent:
    """Container for extracted document content and metadata."""

    text: str
    structure: Dict[str, Any]
    metadata: Dict[str, Any]
    pages: List[Any] = field(default_factory=list)
    embedded_images: List[Any] = field(default_factory=list)
    confidence: float = 1.0


class DocumentProcessor:
    """Processor for structured documents (PDF, DOCX, DOC).

    Extracts text, structure, metadata, images, and tables from documents.
    Supports OCR for scanned pages.
    """

    supported_extensions = ['.pdf', '.docx', '.doc']

    def __init__(self):
        """Initialize document processor."""
        # Note: PyMuPDF check is done at process time to allow class instantiation
        pass

    def process(self, file_path: Path) -> ExtractedContent:
        """Process a document file and extract all content.

        Args:
            file_path: Path to the document file

        Returns:
            ExtractedContent with text, structure, and metadata

        Raises:
            ProcessingError: If processing fails
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise ProcessingError(f"File not found: {file_path}")

        extension = file_path.suffix.lower()

        try:
            if extension == '.pdf':
                return self._process_pdf(file_path)
            elif extension in ['.docx', '.doc']:
                return self._process_docx(file_path)
            else:
                raise ProcessingError(f"Unsupported extension: {extension}")
        except Exception as e:
            if isinstance(e, ProcessingError):
                raise
            raise ProcessingError(f"Failed to process {file_path}: {str(e)}") from e

    def _process_pdf(self, file_path: Path) -> ExtractedContent:
        """Extract content from PDF file.

        Args:
            file_path: Path to PDF file

        Returns:
            ExtractedContent with extracted data
        """
        if fitz is None:
            raise ImportError("PyMuPDF (fitz) is required for PDF processing")

        try:
            doc = fitz.open(file_path)
        except Exception as e:
            raise ProcessingError(f"Failed to open PDF: {str(e)}") from e

        text_parts = []
        pages = []
        embedded_images = []
        metadata = {}

        # Extract metadata
        pdf_metadata = doc.metadata
        if pdf_metadata:
            metadata = {
                'author': pdf_metadata.get('author', ''),
                'title': pdf_metadata.get('title', ''),
                'subject': pdf_metadata.get('subject', ''),
                'creator': pdf_metadata.get('creator', ''),
                'producer': pdf_metadata.get('producer', ''),
                'creationDate': pdf_metadata.get('creationDate', ''),
            }
            # Clean empty values
            metadata = {k: v for k, v in metadata.items() if v}

        metadata['page_count'] = len(doc)

        # Extract text and images from each page
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_text = page.get_text()
            text_parts.append(page_text)

            pages.append({
                'page_num': page_num + 1,
                'text': page_text,
                'is_ocr': False
            })

            # Extract images
            image_list = page.get_images()
            for img_idx, img in enumerate(image_list):
                xref = img[0]
                embedded_images.append({
                    'page': page_num + 1,
                    'xref': xref,
                    'index': img_idx
                })

        doc.close()

        full_text = '\n\n'.join(text_parts)

        # Calculate confidence (simple heuristic: has text = high confidence)
        confidence = 0.9 if full_text.strip() else 0.5

        return ExtractedContent(
            text=full_text,
            structure={'images': len(embedded_images)},
            metadata=metadata,
            pages=pages,
            embedded_images=embedded_images,
            confidence=confidence
        )

    def _process_docx(self, file_path: Path) -> ExtractedContent:
        """Extract content from DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            ExtractedContent with extracted data
        """
        if Document is None:
            raise ProcessingError("python-docx is required for DOCX processing")

        try:
            doc = Document(file_path)
        except Exception as e:
            raise ProcessingError(f"Failed to open DOCX: {str(e)}") from e

        text_parts = []
        headings = []
        tables_data = []

        # Extract paragraphs and headings
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)

                # Check if it's a heading
                style_name = para.style.name
                if style_name == 'Title':
                    # Title style is used for level 0 headings
                    headings.append({
                        'text': text,
                        'level': 0
                    })
                elif style_name.startswith('Heading'):
                    level = 0
                    match = re.search(r'Heading (\d+)', style_name)
                    if match:
                        level = int(match.group(1))
                    headings.append({
                        'text': text,
                        'level': level
                    })

        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables_data.append({
                'index': table_idx,
                'data': table_data
            })

        # Extract metadata from core properties
        metadata = {}
        try:
            props = doc.core_properties
            if props.author:
                metadata['author'] = props.author
            if props.title:
                metadata['title'] = props.title
            if props.subject:
                metadata['subject'] = props.subject
            if props.created:
                metadata['created'] = props.created.isoformat()
            if props.modified:
                metadata['modified'] = props.modified.isoformat()
        except Exception:
            pass  # Metadata extraction is optional

        # Build structure
        structure = {}
        if headings:
            structure['headings'] = headings
        if tables_data:
            structure['tables'] = tables_data

        full_text = '\n\n'.join(text_parts)

        # Calculate confidence
        confidence = 0.95 if full_text.strip() else 0.5

        return ExtractedContent(
            text=full_text,
            structure=structure,
            metadata=metadata,
            pages=[],
            embedded_images=[],
            confidence=confidence
        )

    def _extract_images(self, file_path: Path) -> List[Any]:
        """Extract embedded images from document.

        Args:
            file_path: Path to document

        Returns:
            List of image data
        """
        raise NotImplementedError("Image extraction not yet implemented")

    def _extract_pdf_text(self, doc: Any) -> List[Any]:
        """Extract text from PDF pages.

        Args:
            doc: PyMuPDF document object

        Returns:
            List of page data with text
        """
        pages = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            pages.append(type('Page', (), {
                'text': text,
                'is_ocr': False,
                'page_num': page_num + 1
            })())
        return pages

    def _handle_scanned_pages(self, pages: List[Any]) -> List[Any]:
        """Handle scanned (image-only) pages with OCR.

        Args:
            pages: List of page objects

        Returns:
            Updated pages with OCR text
        """
        raise NotImplementedError("OCR handling not yet implemented")

    def _ocr_page(self, page: Any) -> str:
        """Perform OCR on a page.

        Args:
            page: Page object

        Returns:
            Extracted text
        """
        raise NotImplementedError("OCR not yet implemented")

    def batch_process(self, file_paths: List[Path]) -> List[ExtractedContent]:
        """Process multiple documents in batch.

        Args:
            file_paths: List of document paths

        Returns:
            List of ExtractedContent results
        """
        raise NotImplementedError("Batch processing not yet implemented")
